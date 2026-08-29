"""DOCX export for Build Resume using python-docx (already a project dependency)."""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


TEMPLATE_COLORS = {
    "classic": (0x1F, 0x29, 0x37),
    "modern": (0x1D, 0x4E, 0xD8),
    "minimal": (0x37, 0x41, 0x51),
    "technical": (0x0F, 0x76, 0x6E),
}


def _rgb(name):
    return RGBColor(*TEMPLATE_COLORS.get(name or "classic", TEMPLATE_COLORS["classic"]))


async def generate_resume_docx(sections: list, template_id: str = "classic") -> bytes:
    data = {sec.get("name"): sec.get("data", {}) for sec in sections if isinstance(sec, dict)}
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    accent = _rgb(template_id)
    personal = data.get("personalInfo", {})
    name = personal.get("fullName", "") or ""
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = accent

    role = personal.get("targetRole", "") or ""
    if role:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(role)
        run.font.size = Pt(11)

    contact = "  |  ".join(x for x in [
        personal.get("email", ""), personal.get("phone", ""), personal.get("location", ""),
        personal.get("linkedIn", ""), personal.get("github", ""), personal.get("portfolio", ""),
    ] if x)
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = accent
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10)

    def normal(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        return p

    summary_text = data.get("summary", {}).get("text", "") or personal.get("summary", "")
    if summary_text:
        heading("Summary")
        normal(summary_text)

    raw_skills = data.get("skills", {})
    items = raw_skills.get("items", []) if isinstance(raw_skills, dict) else raw_skills
    if items:
        heading("Skills")
        normal(", ".join(items))

    def entries(title, section_key, field_map):
        raw = data.get(section_key, {})
        entries_list = raw.get("entries", []) if isinstance(raw, dict) else raw
        if not entries_list:
            return
        heading(title)
        for e in entries_list:
            if not isinstance(e, dict):
                continue
            first = field_map.get("first", [])
            line1 = " | ".join(x for x in [e.get(k, "") or "" for k in first] if x)
            if line1:
                normal(line1, bold=True)
            second = " | ".join(x for x in [e.get(k, "") or "" for k in field_map.get("second", [])] if x)
            if second:
                normal(second)
            if field_map.get("meta"):
                meta = ", ".join(x for x in [e.get(k, "") or "" for k in field_map["meta"]] if x)
                if meta:
                    normal(meta)
            desc = e.get("description", "")
            if desc:
                for line in [l.strip() for l in desc.split("\n") if l.strip()]:
                    bullet(line)
            techs = e.get("technology_list") or e.get("techStack", []) or e.get("technologies", [])
            if techs:
                normal("Technologies: " + ", ".join(techs))

    entries("Projects", "projects", {
        "first": ["title"],
        "second": [],
        "meta": [],
    })
    entries("Experience", "experience", {
        "first": ["role", "company"],
        "second": ["location"],
        "meta": ["duration"],
    })
    entries("Internships", "internships", {
        "first": ["role", "company"],
        "second": ["location"],
        "meta": ["duration"],
    })
    entries("Education", "education", {
        "first": ["degree", "branch"],
        "second": ["institute"],
        "meta": ["year", "gpa"],
    })
    entries("Certifications", "certifications", {
        "first": ["name", "issuer"],
        "second": [],
        "meta": ["date"],
    })

    for section_key, title in (("achievements", "Achievements"), ("languages", "Languages")):
        raw = data.get(section_key, {})
        liste = raw.get("items", []) if isinstance(raw, dict) else raw
        if not liste:
            continue
        heading(title)
        for it in liste:
            if isinstance(it, dict):
                bullet(it.get("text", "") or it.get("language", "") or "")
            else:
                bullet(str(it))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()