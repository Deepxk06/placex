import re
import unicodedata
from typing import Dict, List, Optional

from app.services.resume_skills import (
    match_skills_in_text,
    extract_skills_from_section,
    normalize_skill_list,
    skill_occurrence_count,
)
from app.services.resume_sections import detect_sections, get_section_text

SKILL_PATTERNS = {
    "programming": ["python", "java", "javascript", "c\\+\\+", "typescript", "go", "rust", "swift", "kotlin"],
    "web": ["react", "angular", "vue", "node", "django", "flask", "html", "css", "express"],
    "database": ["sql", "mongodb", "postgresql", "mysql", "redis", "oracle", "firebase"],
    "cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins"],
    "data": ["machine learning", "deep learning", "nlp", "tensorflow", "pytorch", "pandas", "numpy"],
    "tools": ["git", "linux", "agile", "scrum", "jira", "ci/cd"],
}

EDUCATION_KEYWORDS = [
    "bachelor", "master", "phd", "b.tech", "m.tech", "b.e", "m.e",
    "b.sc", "m.sc", "bca", "mca", "bba", "mba", "degree", "diploma",
    "university", "college", "institute", "school", "doctorate",
]

BRANCH_PATTERNS = {
    "Computer Science": [r"\bcse\b", r"computer\s*science", r"computer\s*science\s*&?\s*engineering", r"computerscience"],
    "Information Technology": [r"\bit\b(?!\s*(service|company|sector|consultant))", r"information\s*technology", r"information\s*and\s*communication"],
    "Electronics & Communication": [r"\bece\b", r"electronics\s*(and|&)?\s*communication", r"electronic\s*and\s*communication"],
    "Electronics": [r"\be&tc\b", r"electronics"],
    "Electrical": [r"\bee\b", r"electrical"],
    "Mechanical": [r"\bmech(anical)?\b", r"mechanical"],
    "Civil": [r"\bcivil\b"],
    "Data Science": [r"data\s*science", r"\bds\b"],
    "Artificial Intelligence": [r"artificial\s*intelligence", r"\bai\b(?!\s*(tool|based|model))", r"machine\s*learning"],
    "Mathematics": [r"mathematics", r"maths?"],
    "Statistics": [r"statistics"],
    "Physics": [r"physics"],
    "Chemistry": [r"chemistry"],
    "Biotechnology": [r"biotechnology", r"bio\s*technology"],
    "Commerce": [r"commerce"],
    "Management": [r"management", r"\bbba\b", r"\bmba\b"],
    "Arts": [r"arts?\b"],
}

DEGREE_ALIASES = {
    "B.Tech": [r"\bb\.?\s?tech\b", r"bachelor\s*of\s*technology"],
    "M.Tech": [r"\bm\.?\s?tech\b", r"master\s*of\s*technology"],
    "B.E.": [r"\bb\.?\s?e\.?\b(?!\s*tech)", r"bachelor\s*of\s*engineering"],
    "M.E.": [r"\bm\.?\s?e\.?\b(?!\s*tech)", r"master\s*of\s*engineering"],
    "B.Sc": [r"\bb\.?\s?sc\.?\b", r"bachelor\s*of\s*science"],
    "M.Sc": [r"\bm\.?\s?sc\.?\b", r"master\s*of\s*science"],
    "BCA": [r"\bbca\b", r"bachelor\s*of\s*computer\s*applications"],
    "MCA": [r"\bmca\b", r"master\s*of\s*computer\s*applications"],
    "BBA": [r"\bbba\b", r"bachelor\s*of\s*business\s*administration"],
    "MBA": [r"\bmba\b", r"master\s*of\s*business\s*administration"],
    "B.Com": [r"\bb\.?\s?com\.?\b", r"bachelor\s*of\s*commerce"],
    "M.Com": [r"\bm\.?\s?com\.?\b", r"master\s*of\s*commerce"],
    "B.A.": [r"\bb\.?\s?a\.?\b(?!\w)", r"bachelor\s*of\s*arts"],
    "M.A.": [r"\bm\.?\s?a\.?\b(?!\w)", r"master\s*of\s*arts"],
    "Ph.D.": [r"ph\.?\s?d\b", r"doctorate", r"doctor\s*of\s*philosophy"],
    "Diploma": [r"diploma"],
    "Polytechnic": [r"polytechnic"],
    "12th/HSC": [r"higher\s*secondary", r"\bhsc\b", r"12th\b", r"intermediate", r"\+2\b", r"\b12\s*grade\b"],
    "10th/SSC": [r"\bssc\b", r"secondary\s*school", r"10th\b", r"\b10\s*grade\b", r"matriculation"],
}

DATE_RANGE_PATTERNS = [
    r"(\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*\d{4})\s*(?:-|–|to)\s*(\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*\d{4}|present|current|now|till\s*date)",
    r"((?:19|20)\d{2})\s*(?:-|–|to)\s*((?:19|20)\d{2}|present|current|now)",
    r"\(((?:19|20)\d{2})\s*(?:-|–|to)\s*((?:19|20)\d{2}|present)\)",
]

ACTION_VERBS = [
    "achieved", "implemented", "developed", "designed", "optimized", "managed",
    "led", "created", "built", "improved", "delivered", "launched", "engineered",
    "deployed", "integrated", "automated", "configured", "reduced", "increased",
    "established", "generated", "spearheaded", "transformed", "architected",
    "mentored", "coordinated", "facilitated", "streamlined", "accelerated",
]

LOCATION_PATTERN = r"(?:^|\s)([A-Za-z]+(?:[\s-][A-Za-z]+)?,\s*[A-Za-z]{2,}\s*,?\s*(?:India|UK|USA|United\s*States|Canada|Australia|Germany|Singapore|UAE|United\s*Kingdom)?)"

# Unusual symbols that hint at hard-to-parse formatting
UNUSUAL_SYMBOL_PATTERN = re.compile(r"[^\x00-\x7F\u0900-\u097F\u0E00-\u0E7F\u0600-\u06FF\u4E00-\u9FFF\u00C0-\u017F\u2010-\u2027\u20B9\u20AC\u00A3\u00A5\u0020-\u007E]", re.UNICODE)


def normalize_text(text: str) -> str:
    """Normalize Unicode, collapse whitespace, strip zero-width chars."""
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u200b", "").replace("\ufeff", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


async def parse_resume_file(content: bytes, filename: str) -> Dict:
    """Parse PDF / DOCX / TXT content into structured resume data."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            text = await _extract_pdf(content)
        elif ext == "docx":
            text = _extract_docx(content)
        elif ext == "txt":
            text = _extract_txt(content)
        else:
            raise ValueError(f"Unsupported file type: {ext or 'unknown'}")
    except ValueError:
        raise
    except Exception:
        raise ValueError("Could not parse the file. It may be corrupted or not a valid document.")
    if not text or not text.strip():
        raise ValueError("Could not extract any text from the file. The document may be empty or image-only.")
    return extract_resume_data(text)


async def parse_resume_pdf(content: bytes, filename: str) -> Dict:
    """Backward-compatible PDF parser (kept for existing callers)."""
    try:
        text = await _extract_pdf(content)
    except Exception:
        text = content.decode("utf-8", errors="ignore")
    return extract_resume_data(text)


async def _extract_pdf(content: bytes) -> str:
    import fitz  # PyMuPDF (lazy-loaded to keep startup fast)
    doc = fitz.open(stream=content, filetype="pdf")
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return normalize_text("\n".join(parts))


def _extract_docx(content: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text("\n".join(parts))


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return normalize_text(content.decode(encoding))
        except (UnicodeDecodeError, ValueError):
            continue
    return normalize_text(content.decode("utf-8", errors="ignore"))


def clean_text(text: str) -> str:
    return normalize_text(text)


def extract_resume_data(text: str) -> Dict:
    text = clean_text(text)
    if not text:
        return {}
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    sections = detect_sections(text)
    email = extract_email(text)
    phone = extract_phone(text)
    linkedin, github, portfolio = extract_links(text)
    name = extract_name(lines, sections)
    summary = extract_summary(sections)
    target_role = extract_target_role(sections, text)
    location = extract_location(lines)

    skills_section_text = get_section_text(sections, "skills") + "\n" + get_section_text(sections, "soft_skills")
    skills_list = extract_skills_from_section(skills_section_text)
    skills_in_text = match_skills_in_text(text)
    all_skills = normalize_skill_list(list(dict.fromkeys(skills_list + skills_in_text)))
    skill_details = _build_skill_details(all_skills, skills_list, text)

    education = extract_education(text, sections)
    experience = extract_experience(text, sections)
    internships = extract_internships(text, sections)
    projects = extract_projects(text, sections)
    certifications = extract_certifications(text, sections)
    achievements = extract_achievements(sections)
    languages = extract_languages(sections)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "location": location,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
        "targetRole": target_role,
        "summary": summary,
        "skills": all_skills,
        "skillDetails": skill_details,
        "education": education,
        "experience": experience,
        "internships": internships,
        "projects": projects,
        "certifications": certifications,
        "achievements": achievements,
        "languages": languages,
        "detectedSections": list(sections.keys()),
        "rawText": text,
    }


def _build_skill_details(all_skills: List[str], listed_skills: List[str], text: str) -> Dict[str, Dict]:
    listed = set(s.lower() for s in listed_skills)
    details = {}
    for skill in all_skills:
        occurrences = skill_occurrence_count(text, skill)
        if skill.lower() in listed:
            level = "strong"
        elif occurrences >= 2:
            level = "mentioned"
        else:
            level = "weak"
        details[skill] = {"level": level, "occurrences": occurrences}
    return details


def extract_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    return match.group(0) if match else ""


def extract_links(text: str) -> tuple:
    linkedin = re.search(r"linkedin\.com/[^\s\]+]+", text)
    github = re.search(r"github\.com/[A-Za-z0-9_.-]+", text)
    portfolio = None
    for match in re.finditer(r"https?://[^\s]+", text):
        url = match.group(0).rstrip(".,;)]")
        if "linkedin" in url.lower() or "github" in url.lower():
            continue
        portfolio = url
        break
    return (linkedin.group(0) if linkedin else "",
            github.group(0) if github else "",
            portfolio or "")


def extract_name(lines: List[str], sections: Dict) -> str:
    heading_names = set()
    for data in sections.values():
        heading_names.add(str(data.get("heading", "")).lower())
    for line in lines:
        if len(line) > 60:
            continue
        if line.lower() in heading_names:
            continue
        if "@" in line or re.search(r"\d{3}[-.\s]?\d{3}", line) or "http" in line.lower():
            continue
        if not any(ch.isdigit() for ch in line):
            words = line.split()
            if 1 <= len(words) <= 5:
                return line
    return ""


def extract_summary(sections: Dict) -> str:
    body = get_section_text(sections, "summary")
    if not body:
        body = get_section_text(sections, "objective")
    return body[:1000]


def extract_target_role(sections: Dict, text: str) -> str:
    summary = get_section_text(sections, "summary") + "\n" + get_section_text(sections, "objective")
    lines = [l.strip() for l in summary.split("\n") if l.strip()][:3]
    for line in lines:
        match = re.search(r"(?:seeking|targeting|for|as)\s+a[n]?\s+([A-Z][A-Za-z ]{2,40}?(?:Developer|Engineer|Scientist|Analyst|Designer|Manager|Consultant|Intern|Lead|Architect|Trainee))", line, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    match = re.search(r"\b((?:Software|Data|ML|AI|Machine Learning|Frontend|Backend|Full[- ]Stack|DevOps|Cloud|Product|UI|UX|Systems|Security|Quality|Test)[ ]?(?:Engineer|Developer|Scientist|Analyst|Architect|Trainee))\b", text, re.IGNORECASE)
    return match.group(1) if match else ""


def extract_location(lines: List[str]) -> str:
    for line in lines:
        if "@" in line or "http" in line.lower() or re.search(r"\b\d{5,6}\b", line):
            match = re.search(r"([A-Za-z]+(?:[\s-][A-Za-z]+)?)[, ]+([A-Za-z]{2,})[, ]*(India)?", line)
            if match:
                return f"{match.group(1)}, {match.group(2)}" + (f", {match.group(3)}" if match.group(3) else "")
    for line in lines:
        if re.search(r"\b\d{6}\b", line) and len(line) < 40:
            return line.strip()
    return ""


def extract_skills(text: str) -> list:
    found = set()
    text_lower = text.lower()
    for category, patterns in SKILL_PATTERNS.items():
        for pattern in patterns:
            if re.search(r"\b" + pattern + r"\b", text_lower):
                found.add(pattern.replace("\\", "").replace("+", "+"))
    return sorted(found)


def extract_education(text: str, sections: Dict) -> List[Dict]:
    body = get_section_text(sections, "education")
    source = body if body else text
    lines = [l.strip() for l in source.split("\n") if l.strip()]
    education: List[Dict] = []
    current: Dict = {}
    for line in lines:
        lower = line.lower()
        degree = _detect_degree(line)
        is_entry = degree or any(kw in lower for kw in EDUCATION_KEYWORDS)
        if is_entry:
            if current:
                education.append(current)
            current = {
                "degree": degree or line,
                "institute": "",
                "year": "",
                "gpa": "",
                "branch": "",
            }
            year = _detect_year(line)
            if year:
                current["year"] = year
            gpa = _detect_gpa(line)
            if gpa:
                current["gpa"] = gpa
            branch = _detect_branch(line)
            if branch:
                current["branch"] = branch
            institute = _detect_institute(line, degree)
            if institute:
                current["institute"] = institute
        elif current and len(current.get("institute", "")) < 80:
            year = _detect_year(line)
            gpa = _detect_gpa(line)
            branch = _detect_branch(line)
            if year and not current.get("year"):
                current["year"] = year
            elif not current.get("institute"):
                current["institute"] = line
            if gpa and not current.get("gpa"):
                current["gpa"] = gpa
            if branch and not current.get("branch"):
                current["branch"] = branch
            elif current.get("institute") and not year and not gpa and len(line) < 80:
                if not _detect_degree(line):
                    current["institute"] = line
    if current:
        education.append(current)
    return education[:6]


def _detect_degree(line: str) -> str:
    lower = line.lower()
    for degree, patterns in DEGREE_ALIASES.items():
        for pattern in patterns:
            if re.search(pattern, lower):
                return degree
    return ""


def _detect_branch(line: str) -> str:
    lower = line.lower()
    for branch, patterns in BRANCH_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, lower):
                return branch
    return ""


def _detect_institute(line: str, degree: str) -> str:
    """Best-effort extraction of institute name from an education line."""
    if not degree:
        return ""
    rest = line
    if degree and degree.lower() not in rest.lower():
        return ""
    # strip the degree token(s)
    for pattern in DEGREE_ALIASES.get(degree, []):
        rest = re.sub(pattern, " ", rest, flags=re.IGNORECASE)
    rest = re.sub(r"\b((?:19|20)\d{2})\b", " ", rest)
    rest = re.sub(r"(?:cgpa|gpa|percentage|score)\s*[:=\-]?\s*\d+(?:\.\d{1,2})?\s*(?:/\s*\d+(?:\.\d+)?|%)?", " ", rest, flags=re.IGNORECASE)
    for branch, patterns in BRANCH_PATTERNS.items():
        for pattern in patterns:
            rest = re.sub(pattern, " ", rest, flags=re.IGNORECASE)
    rest = rest.replace(",", " ")
    rest = re.sub(r"\s{2,}", " ", rest).strip()
    if len(rest) < 4 or len(rest) > 60:
        return ""
    if any(kw in rest.lower() for kw in ("university", "college", "institute", "school", "academy")):
        return rest
    return rest


def _detect_year(line: str) -> str:
    match = re.search(r"\b((?:19|20)\d{2})\b", line)
    return match.group(1) if match else ""


def _detect_gpa(line: str) -> str:
    match = re.search(r"(?:cgpa|gpa|percentage|score)\s*[:=\-]?\s*(\d+(?:\.\d{1,2})?)\s*(?:/\s*\d+(?:\.\d+)?|%)?", line, re.IGNORECASE)
    if match:
        return match.group(1)
    match = re.search(r"(\d+\.\d{1,2})\s*(?:/\s*10|out\s*of\s*10|/4)", line, re.IGNORECASE)
    return match.group(1) if match else ""


def _extract_duration(line: str) -> str:
    for pattern in DATE_RANGE_PATTERNS:
        match = re.search(pattern, line, re.IGNORECASE)
        if match:
            return match.group(0).strip("()")
    return ""


def extract_experience(text: str, sections: Dict) -> List[Dict]:
    body = get_section_text(sections, "experience")
    if not body:
        return []
    lines = [l.strip() for l in body.split("\n") if l.strip()]
    entries: List[Dict] = []
    current: Dict = {}
    for line in lines:
        duration = _extract_duration(line)
        lower = line.lower()
        if duration and (current.get("role") or current.get("company")):
            current["duration"] = duration
            continue
        is_company_like = re.match(r"^[A-Za-z][A-Za-z0-9&\-. ]{1,40}$", line) and not any(v in lower for v in ACTION_VERBS[:6]) and len(line) < 45
        if not current.get("role") and not current.get("company"):
            current = {"company": line, "role": "", "duration": "", "description": "", "technologies": []}
            dur = _extract_duration(line)
            if dur:
                current["duration"] = dur
        elif current.get("role") is None:
            current["role"] = line
        elif not current.get("role"):
            current["role"] = line
        elif len(line) < 45 and not any(v in lower for v in ACTION_VERBS) and not _contains_punctuation_sentence(line):
            current["role"] = f"{current.get('role', '')} / {line}".strip(" /")
        else:
            current["description"] = (current.get("description", "") + " " + line).strip()
        if not current.get("company") and current.get("role"):
            current["company"] = ""
    if current:
        current["technologies"] = normalize_skill_list(match_skills_in_text(current.get("description", "")))
        entries.append(current)
    for entry in entries:
        if not entry.get("technologies"):
            entry["technologies"] = normalize_skill_list(match_skills_in_text(entry.get("description", "")))
    return entries


def _contains_punctuation_sentence(line: str) -> bool:
    return bool(re.search(r"[.!?]\s*$", line)) or len(line) > 60


def extract_internships(text: str, sections: Dict) -> List[Dict]:
    body = get_section_text(sections, "internships")
    if not body:
        return []
    lines = [l.strip() for l in body.split("\n") if l.strip()]
    entries: List[Dict] = []
    current: Dict = {}
    for line in lines:
        duration = _extract_duration(line)
        if duration and (current.get("role") or current.get("company")):
            current["duration"] = duration
            continue
        if not current.get("role") and not current.get("company"):
            current = {"company": line, "role": "", "duration": duration, "description": ""}
        elif not current.get("role"):
            current["role"] = line
        else:
            current["description"] = (current.get("description", "") + " " + line).strip()
    if current:
        current["technologies"] = normalize_skill_list(match_skills_in_text(current.get("description", "")))
        entries.append(current)
    for entry in entries:
        if not entry.get("technologies"):
            entry["technologies"] = normalize_skill_list(match_skills_in_text(entry.get("description", "")))
    return entries


def extract_projects(text: str, sections: Dict) -> List[Dict]:
    body = get_section_text(sections, "projects")
    if not body:
        return []
    lines = [l.strip() for l in body.split("\n") if l.strip()]
    projects: List[Dict] = []
    current: Dict = {}
    for line in lines:
        lower = line.lower()
        tech = match_skills_in_text(line)
        if not current or _looks_like_project_title(line, lower):
            if current:
                projects.append(current)
            current = {"title": line, "description": "", "techStack": [], "link": ""}
            link = re.search(r"https?://[^\s]+", line)
            if link:
                current["link"] = link.group(0).rstrip(".,;)]")
        else:
            current["description"] = (current.get("description", "") + " " + line).strip()
        new_tech = [t for t in tech if t not in current["techStack"]]
        current["techStack"].extend(new_tech)
        link = re.search(r"https?://[^\s]+", line)
        if link and not current.get("link"):
            current["link"] = link.group(0).rstrip(".,;)]")
    if current:
        projects.append(current)
    return projects[:8]


def _looks_like_project_title(line: str, lower: str) -> bool:
    if len(line) > 70:
        return False
    if not line[0].isalpha():
        return False
    if line.count(" ") <= 6 and " and " not in lower and "the " not in lower and "with " not in lower:
        if not lower.endswith((".", "!", "?")):
            return True
    return bool(re.match(r"^[A-Z][A-Za-z0-9'&.,\s\-]{3,35}$", line)) and not lower.endswith(".")


def extract_certifications(text: str, sections: Dict) -> List[Dict]:
    body = get_section_text(sections, "certifications")
    source = body if body else text
    lines = [l.strip() for l in source.split("\n") if l.strip()]
    certs: List[Dict] = []
    for i, line in enumerate(lines):
        lower = line.lower()
        if any(kw in lower for kw in ["certification", "certificate", "certified", "credential", "coursera", "udemy", "nptel"]) and len(line) < 80:
            entry = {"name": line.strip(" ."), "issuer": ""}
            if i + 1 < len(lines) and len(lines[i + 1]) < 60:
                next_lower = lines[i + 1].lower()
                if not any(kw in next_lower for kw in ["certification", "certificate", "certified"]):
                    entry["issuer"] = lines[i + 1].strip()
            certs.append(entry)
    return certs[:10]


def extract_achievements(sections: Dict) -> List[str]:
    body = get_section_text(sections, "achievements")
    if not body:
        return []
    lines = [l.strip().lstrip("-*•·›") for l in body.split("\n") if l.strip()]
    return [l for l in lines if l][:10]


def extract_languages(sections: Dict) -> List[str]:
    body = get_section_text(sections, "languages")
    if not body:
        return []
    lines = [l.strip().lstrip("-*•·›") for l in body.split("\n") if l.strip()]
    langs = []
    for line in lines:
        parts = re.split(r"[,|;]", line)
        for part in parts:
            part = part.strip()
            if part and re.match(r"^[A-Za-z][A-Za-z\s'.\-]+$", part):
                langs.append(re.sub(r"\s*\([^)]*\)", "", part).strip())
    return langs[:10]
